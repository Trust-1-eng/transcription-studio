import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

import io
from docx import Document
from app.importers.text_extract import extract_text_from_docx, extract_text_from_pdf, extract_text


def _make_docx_bytes(paragraphs: list) -> bytes:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_docx():
    text = "The quick brown fox jumps over the lazy dog and keeps running forward."
    data = _make_docx_bytes([text])
    result = extract_text_from_docx(data)
    assert "quick brown fox" in result


def test_extract_docx_multiple_paragraphs():
    paras = [
        "First paragraph with enough words to pass validation.",
        "Second paragraph also with enough words to pass.",
        "Third paragraph for good measure and testing purposes.",
    ]
    data = _make_docx_bytes(paras)
    result = extract_text_from_docx(data)
    assert "First paragraph" in result
    assert "Second paragraph" in result


def test_extract_text_routing_docx():
    text = "This is a test document with more than ten words for validation purposes."
    data = _make_docx_bytes([text])
    result = extract_text("test.docx", data)
    assert "test document" in result


def test_extract_text_unsupported():
    try:
        extract_text("file.xyz", b"data")
        assert False, "Should have raised ValueError"
    except ValueError as e:
        assert "Unsupported" in str(e)


def test_extract_text_too_short():
    data = _make_docx_bytes(["short"])
    try:
        extract_text("test.docx", data)
        assert False, "Should have raised ValueError"
    except ValueError as e:
        assert "meaningful text" in str(e)


def _make_sample_docx(path: str):
    """Create a sample DOCX file for manual testing."""
    doc = Document()
    doc.add_paragraph("The quick brown fox jumps over the lazy dog. This is a sample edited transcript for alignment testing.")
    doc.add_paragraph("Dr. Johnson spoke about the importance of education. He mentioned that every student deserves quality instruction.")
    doc.add_paragraph("The meeting concluded at three pm. All participants agreed on the next steps going forward.")
    doc.save(path)
    print(f"Sample DOCX created: {path}")


if __name__ == "__main__":
    test_extract_docx()
    test_extract_docx_multiple_paragraphs()
    test_extract_text_routing_docx()
    test_extract_text_unsupported()
    test_extract_text_too_short()
    print("All extract tests passed!")

    # Create sample file for manual testing
    samples_dir = Path(__file__).parent / "samples"
    samples_dir.mkdir(exist_ok=True)
    _make_sample_docx(str(samples_dir / "sample.docx"))
