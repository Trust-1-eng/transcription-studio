import io

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from fpdf import FPDF

from app.config import DEJAVU_PATH, DEJAVU_BOLD_PATH
from app.utils import ms_to_readable
from app.dependencies import gemini_cache
from app.exporters.helpers import make_doc_header, resolve_speakers, HEADER_DEFAULT


def _set_cell_shading(cell, color_hex: str):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def make_table_docx(data: dict, label_map: dict = None, speaker_count: int = None) -> bytes:
    tid = data.get("id", "")
    gemini = gemini_cache.get(tid)
    gdata = gemini[0] if gemini and not gemini[0].get("error") else {}
    if label_map is None:
        label_map, speaker_count = resolve_speakers(data, gdata)

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    doc.add_heading(gdata.get("title", "Transcript"), 0)

    hdr = make_doc_header(data)
    if speaker_count is not None:
        hdr["speakers"] = speaker_count
    lb = hdr["labels"]
    meta = f"{lb['date']}: {hdr['date']}    {lb['duration']}: {hdr['duration']}    {lb['lang']}: {hdr['language']}"
    if hdr["speakers"]:
        meta += f"    {lb['speakers']}: {hdr['speakers']}"
    p = doc.add_paragraph()
    run = p.add_run(meta)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    utterances = data.get("utterances", [])
    if not utterances:
        doc.add_paragraph(data.get("text", ""))
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # Fixed table layout to prevent Word from auto-adjusting column widths
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    # Header row
    headers = ["№", "Taймкод", "Спікер", "Текст"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(cell, "2D2D2D")

    # Data rows
    for idx, u in enumerate(utterances, 1):
        row = table.add_row()
        ts = ms_to_readable(u.get("start", 0))
        spk = label_map.get(u.get("speaker", ""), u.get("speaker", "Speaker"))
        text = u.get("text", "")

        values = [str(idx), ts, spk, text]
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)

        # Alternating row shading
        if idx % 2 == 0:
            for cell in row.cells:
                _set_cell_shading(cell, "F5F5F5")

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(0.8)
        row.cells[1].width = Cm(2.0)
        row.cells[2].width = Cm(3.0)
        row.cells[3].width = Cm(12.2)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_table_pdf(data: dict, label_map: dict = None, speaker_count: int = None) -> bytes:
    tid = data.get("id", "")
    gemini = gemini_cache.get(tid)
    gdata = gemini[0] if gemini and not gemini[0].get("error") else {}
    if label_map is None:
        label_map, speaker_count = resolve_speakers(data, gdata)

    use_unicode = DEJAVU_PATH.exists()
    use_bold = DEJAVU_BOLD_PATH.exists()
    _font = "DejaVu" if use_unicode else "Helvetica"

    title = gdata.get("title", "Transcript")

    class TablePDF(FPDF):
        def header(self):
            if self.page_no() == 1:
                self.set_font(_font, size=14)
                self.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT", align="C")
                hdr_info = make_doc_header(data)
                if speaker_count is not None:
                    hdr_info["speakers"] = speaker_count
                lb = hdr_info.get("labels", HEADER_DEFAULT)
                self.set_font(_font, size=8)
                self.set_text_color(120, 120, 120)
                meta = f"{lb['date']}: {hdr_info['date']}  |  {lb['duration']}: {hdr_info['duration']}  |  {lb['lang']}: {hdr_info['language']}"
                if hdr_info.get("speakers"):
                    meta += f"  |  {lb['speakers']}: {hdr_info['speakers']}"
                self.cell(0, 6, meta, new_x="LMARGIN", new_y="NEXT", align="C")
                self.set_text_color(0, 0, 0)
                self.ln(4)

        def footer(self):
            self.set_y(-15)
            self.set_font(_font, size=8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 10, f"{self.page_no()} / {{nb}}", align="C")

    pdf = TablePDF()
    pdf.alias_nb_pages()
    if use_unicode:
        pdf.add_font("DejaVu", fname=str(DEJAVU_PATH))
        if use_bold:
            pdf.add_font("DejaVu", style="B", fname=str(DEJAVU_BOLD_PATH))
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    utterances = data.get("utterances", [])
    if not utterances:
        pdf.set_font(_font, size=11)
        pdf.multi_cell(0, 5, data.get("text", ""))
        buf = io.BytesIO()
        pdf.output(buf)
        return buf.getvalue()

    page_w = pdf.w - pdf.l_margin - pdf.r_margin
    col_w = [10, 22, 35, page_w - 67]  # №, Таймкод, Спікер, Текст

    # Table header
    pdf.set_fill_color(45, 45, 45)
    pdf.set_text_color(255, 255, 255)
    if use_bold:
        pdf.set_font(_font, style="B", size=9)
    else:
        pdf.set_font(_font, size=9)

    headers = ["№", "Таймкод", "Спікер", "Текст"]
    for i, h in enumerate(headers):
        pdf.cell(col_w[i], 7, h, border=1, fill=True)
    pdf.ln()

    pdf.set_text_color(0, 0, 0)
    pdf.set_font(_font, size=8)

    for idx, u in enumerate(utterances, 1):
        ts = ms_to_readable(u.get("start", 0))
        spk = label_map.get(u.get("speaker", ""), u.get("speaker", "Speaker"))
        text = u.get("text", "")

        # Calculate row height based on text
        text_lines = pdf.multi_cell(col_w[3], 4, text, dry_run=True, output="LINES")
        row_h = max(5, len(text_lines) * 4 + 1)

        # Check page break
        if pdf.get_y() + row_h > pdf.h - 25:
            pdf.add_page()
            # Repeat header on new page
            pdf.set_fill_color(45, 45, 45)
            pdf.set_text_color(255, 255, 255)
            if use_bold:
                pdf.set_font(_font, style="B", size=9)
            else:
                pdf.set_font(_font, size=9)
            for i, h in enumerate(headers):
                pdf.cell(col_w[i], 7, h, border=1, fill=True)
            pdf.ln()
            pdf.set_text_color(0, 0, 0)
            pdf.set_font(_font, size=8)

        # Alternating row color
        if idx % 2 == 0:
            pdf.set_fill_color(245, 245, 245)
            fill = True
        else:
            fill = False

        y_start = pdf.get_y()
        x_start = pdf.get_x()

        # Draw cells
        pdf.cell(col_w[0], row_h, str(idx), border=1, fill=fill)
        pdf.cell(col_w[1], row_h, ts, border=1, fill=fill)
        pdf.cell(col_w[2], row_h, spk, border=1, fill=fill)

        # Text cell with wrapping
        x_text = pdf.get_x()
        if fill:
            pdf.rect(x_text, y_start, col_w[3], row_h, style="F")
        pdf.rect(x_text, y_start, col_w[3], row_h)
        pdf.set_xy(x_text + 1, y_start + 0.5)
        pdf.multi_cell(col_w[3] - 2, 4, text)
        pdf.set_xy(x_start, y_start + row_h)

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()
