import io

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

from app.utils import ms_to_readable
from app.dependencies import gemini_cache
from app.exporters.helpers import make_doc_header, resolve_speakers


def make_docx(data: dict, title: str = "Transcript", label_map: dict = None, speaker_count: int = None) -> bytes:
    tid = data.get("id", "")
    gemini = gemini_cache.get(tid)
    gdata = gemini[0] if gemini and not gemini[0].get("error") else {}
    if label_map is None:
        label_map, speaker_count = resolve_speakers(data, gdata)

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(4)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    doc.add_heading(gdata.get("title", title), 0)

    hdr = make_doc_header(data)
    if speaker_count is not None:
        hdr["speakers"] = speaker_count
    lb = hdr["labels"]
    meta = f"{lb['date']}: {hdr['date']}    {lb['duration']}: {hdr['duration']}    {lb['lang']}: {hdr['language']}"
    if hdr["speakers"]:
        meta += f"    {lb['speakers']}: {hdr['speakers']}"
    p = doc.add_paragraph()
    run = p.add_run(meta)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    if gdata.get("summary"):
        p = doc.add_paragraph()
        run = p.add_run(gdata["summary"])
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Thin horizontal rule separator
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(4)
    sep.paragraph_format.space_after = Pt(8)
    pPr = sep._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

    if data.get("utterances"):
        for u in data["utterances"]:
            ts = ms_to_readable(u.get("start", 0))
            spk = label_map.get(u.get("speaker", ""), u.get("speaker", "Speaker"))
            p = doc.add_paragraph()
            run_label = p.add_run(f"[{ts}] {spk}: ")
            run_label.bold = True
            p.add_run(u.get("text", ""))
    else:
        doc.add_paragraph(data.get("text", ""))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_bilingual_docx(data: dict, translation: dict, label_map: dict = None, speaker_count: int = None) -> bytes:
    tid = data.get("id", "")
    gemini = gemini_cache.get(tid)
    gdata = gemini[0] if gemini and not gemini[0].get("error") else {}
    if label_map is None:
        label_map, speaker_count = resolve_speakers(data, gdata)

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(2)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    doc.add_heading(gdata.get("title", "Bilingual Transcript"), 0)

    orig_utts = data.get("utterances", [])
    trans_utts = translation.get("utterances", [])

    if orig_utts and trans_utts:
        for i, u in enumerate(orig_utts):
            ts = ms_to_readable(u.get("start", 0))
            spk = label_map.get(u.get("speaker", ""), u.get("speaker", "Speaker"))

            p = doc.add_paragraph()
            run_label = p.add_run(f"[{ts}] {spk}: ")
            run_label.bold = True
            p.add_run(u.get("text", ""))

            if i < len(trans_utts):
                p2 = doc.add_paragraph()
                run_arrow = p2.add_run("\u2192 ")
                run_arrow.font.color.rgb = RGBColor(0x66, 0x99, 0xCC)
                run_trans = p2.add_run(trans_utts[i].get("text", ""))
                run_trans.italic = True
                run_trans.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            # Small spacing between pairs
            sep = doc.add_paragraph()
            sep.paragraph_format.space_before = Pt(0)
            sep.paragraph_format.space_after = Pt(4)
    elif orig_utts:
        doc.add_heading("Оригінал", level=1)
        for u in orig_utts:
            ts = ms_to_readable(u.get("start", 0))
            spk = label_map.get(u.get("speaker", ""), u.get("speaker", "Speaker"))
            p = doc.add_paragraph()
            run_label = p.add_run(f"[{ts}] {spk}: ")
            run_label.bold = True
            p.add_run(u.get("text", ""))
        doc.add_paragraph()
        doc.add_heading("Переклад", level=1)
        trans_text = translation.get("text", "")
        if trans_text:
            doc.add_paragraph(trans_text)
        else:
            p = doc.add_paragraph()
            run = p.add_run("(Текст перекладу недоступний)")
            run.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    else:
        doc.add_heading("Оригінал", level=1)
        doc.add_paragraph(data.get("text", ""))
        doc.add_paragraph()
        doc.add_heading("Переклад", level=1)
        trans_text = translation.get("text", "")
        if trans_text:
            doc.add_paragraph(trans_text)
        else:
            p = doc.add_paragraph()
            run = p.add_run("(Текст перекладу недоступний)")
            run.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


PAUSE_THRESHOLD_MS = 3000
LONG_PAUSE_THRESHOLD_MS = 8000


def make_verbatim_docx(data: dict, label_map: dict = None, speaker_count: int = None) -> bytes:
    tid = data.get("id", "")
    gemini = gemini_cache.get(tid)
    gdata = gemini[0] if gemini and not gemini[0].get("error") else {}
    if label_map is None:
        label_map, speaker_count = resolve_speakers(data, gdata)

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(4)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    doc.add_heading(gdata.get("title", "Verbatim Transcript"), 0)

    hdr = make_doc_header(data)
    if speaker_count is not None:
        hdr["speakers"] = speaker_count
    lb = hdr["labels"]
    meta = f"{lb['date']}: {hdr['date']}    {lb['duration']}: {hdr['duration']}    {lb['lang']}: {hdr['language']}"
    if hdr["speakers"]:
        meta += f"    {lb['speakers']}: {hdr['speakers']}"
    p = doc.add_paragraph()
    run = p.add_run(meta)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Separator
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(4)
    sep.paragraph_format.space_after = Pt(8)
    pPr = sep._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

    utterances = data.get("utterances", [])
    if not utterances:
        doc.add_paragraph(data.get("text", ""))
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    for i, u in enumerate(utterances):
        ts = ms_to_readable(u.get("start", 0))
        spk = label_map.get(u.get("speaker", ""), u.get("speaker", "Speaker"))
        text = u.get("text", "")

        # Annotation before utterance
        if i > 0:
            prev = utterances[i - 1]
            gap = u.get("start", 0) - prev.get("end", prev.get("start", 0))

            annotation = None
            if gap >= LONG_PAUSE_THRESHOLD_MS:
                annotation = f"(довга пауза \u2014 {gap / 1000:.0f} сек)"
            elif gap >= PAUSE_THRESHOLD_MS:
                annotation = f"(пауза \u2014 {gap / 1000:.0f} сек)"
            elif gap < -200:
                annotation = "(перебив)"

            if annotation:
                p_ann = doc.add_paragraph()
                p_ann.alignment = 1  # center
                run_ann = p_ann.add_run(annotation)
                run_ann.italic = True
                run_ann.font.size = Pt(9)
                run_ann.font.color.rgb = RGBColor(0xCC, 0x66, 0x33)

        p = doc.add_paragraph()
        run_label = p.add_run(f"[{ts}] {spk}: ")
        run_label.bold = True
        p.add_run(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_interview_docx(data: dict, label_map: dict = None, speaker_count: int = None) -> bytes:
    tid = data.get("id", "")
    gemini = gemini_cache.get(tid)
    gdata = gemini[0] if gemini and not gemini[0].get("error") else {}
    if label_map is None:
        label_map, speaker_count = resolve_speakers(data, gdata)

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(4)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    doc.add_heading(gdata.get("title", "Interview"), 0)

    hdr = make_doc_header(data)
    if speaker_count is not None:
        hdr["speakers"] = speaker_count
    lb = hdr["labels"]
    meta = f"{lb['date']}: {hdr['date']}    {lb['duration']}: {hdr['duration']}    {lb['lang']}: {hdr['language']}"
    if hdr["speakers"]:
        meta += f"    {lb['speakers']}: {hdr['speakers']}"
    p = doc.add_paragraph()
    run = p.add_run(meta)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Thin separator
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(4)
    sep.paragraph_format.space_after = Pt(8)
    pPr = sep._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

    utterances = data.get("utterances", [])
    if not utterances:
        doc.add_paragraph(data.get("text", ""))
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # Detect interviewer: first speaker in the transcript
    interviewer = utterances[0].get("speaker", "")

    for u in utterances:
        spk_raw = u.get("speaker", "")
        spk = label_map.get(spk_raw, spk_raw)
        text = u.get("text", "")
        is_interviewer = (spk_raw == interviewer)

        p = doc.add_paragraph()

        if is_interviewer:
            # Interviewer: bold name + bold text
            run_label = p.add_run(f"{spk}: ")
            run_label.bold = True
            run_label.font.size = Pt(11)
            run_text = p.add_run(text)
            run_text.bold = True
            run_text.font.size = Pt(11)
        else:
            # Respondent: normal name, normal text, left indent
            p.paragraph_format.left_indent = Cm(1)
            run_label = p.add_run(f"{spk}: ")
            run_label.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run_label.font.size = Pt(11)
            p.add_run(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_literary_docx(data: dict, label_map: dict = None, speaker_count: int = None) -> bytes:
    tid = data.get("id", "")
    gemini = gemini_cache.get(tid)
    gdata = gemini[0] if gemini and not gemini[0].get("error") else {}
    if label_map is None:
        label_map, speaker_count = resolve_speakers(data, gdata)

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    doc.add_heading(gdata.get("title", "Transcript"), 0)

    if data.get("utterances"):
        prev_speaker = None
        for u in data["utterances"]:
            spk = label_map.get(u.get("speaker", ""), u.get("speaker", ""))
            text = u.get("text", "")
            p = doc.add_paragraph()
            if spk and spk != prev_speaker:
                run_label = p.add_run(f"\u2014 {spk}: ")
                run_label.bold = True
                p.add_run(text)
            else:
                p.add_run(f"\u2014 {text}")
            prev_speaker = spk
    else:
        doc.add_paragraph(data.get("text", ""))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_paragraphs_docx(paragraphs: list) -> bytes:
    doc = Document()
    doc.add_heading("Paragraphs", 0)
    for p_data in paragraphs:
        ts = ms_to_readable(p_data.get("start", 0))
        p = doc.add_paragraph()
        run_ts = p.add_run(f"[{ts}] ")
        run_ts.font.color.rgb = RGBColor(0x88, 0x92, 0xA4)
        p.add_run(p_data.get("text", ""))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
